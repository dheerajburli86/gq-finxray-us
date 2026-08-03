import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt

pdf_path = "ss.pdf"
docx_path = "output.docx"

print("Opening PDF...")
pdf = fitz.open(pdf_path)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

total_pages = len(pdf)

print(f"Total pages: {total_pages}")

for i, page in enumerate(pdf):
    try:
        print(f"Processing page {i+1}/{total_pages}")

        text = page.get_text("text")

        if text.strip():
            doc.add_paragraph(text)
        else:
            doc.add_paragraph("[No selectable text found on this page]")

        doc.add_page_break()

        # Save every 100 pages so progress isn't lost
        if (i + 1) % 100 == 0:
            doc.save(docx_path)
            print(f"Saved progress at page {i+1}")

    except Exception as e:
        print(f"Error on page {i+1}: {e}")

print("Saving final document...")
doc.save(docx_path)

print("Done!")